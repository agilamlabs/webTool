"""Three-tier content extraction: trafilatura -> BeautifulSoup4 -> raw text.

Also supports binary extraction for PDF (pypdf), XLSX (openpyxl), DOCX
(python-docx), and CSV (stdlib). PDF/XLSX/DOCX require the optional
``[binary]`` extra; CSV works with no additional dependency. Without
the relevant library, binary extraction returns
``ExtractionResult(extraction_method="none")`` with a clear install
hint -- it never crashes the pipeline.
"""

from __future__ import annotations

import asyncio
import importlib.util
import json
import re
from typing import Any, Optional
from urllib.parse import urlparse

import trafilatura
from bs4 import BeautifulSoup
from loguru import logger

from .config import AppConfig
from .injection import (
    detect_injection,
    redact_injection,
    strip_hidden_dom,
    strip_invisible_chars,
)
from .models import ExtractionResult, FetchResult, FetchStatus, InjectionReport

# v1.6.16 CE-3: URL fragments that mark a captured response as
# analytics/telemetry rather than the page's real data API. Used to
# DEPRIORITISE such bodies when ``prefer_api`` picks among captured XHR/fetch
# JSON responses, so a large batched analytics payload no longer wins over a
# small, relevant same-origin API response.
_ANALYTICS_URL_RE = re.compile(
    r"(?:analytics|telemetry|tracking|/track\b|/collect\b|beacon|segment"
    r"|google-analytics|googletagmanager|gtag|/gtm\b|doubleclick|/pixel\b"
    r"|/stats\b|/metrics\b|/rum\b|/log(?:s|ging)?\b)",
    re.IGNORECASE,
)


def _module_available(name: str) -> bool:
    """v1.7.0 Wave 3C: True when an optional module can be imported.

    Probes via :func:`importlib.util.find_spec` without importing the
    module, so the no-``[binary]``-extra install is unaffected. Used to
    choose the PDF engine (pdfplumber preferred, pypdf fallback). Tests
    patch this to exercise each branch.
    """
    try:
        return importlib.util.find_spec(name) is not None
    except (ImportError, ValueError):  # pragma: no cover -- defensive
        # ImportError: a parent package is itself broken/missing.
        # ValueError: ``name`` has no spec machinery (e.g. __main__).
        return False


def _pdf_title_from_metadata(metadata: Any) -> Optional[str]:
    """Pull a non-empty ``/Title`` from pypdf/pdfplumber metadata.

    pypdf exposes a ``DocumentInformation`` object (``.title`` attr);
    pdfplumber exposes a plain dict (``{'Title': ...}``). Handle both
    shapes; never raise -- a missing/garbage title just yields None.
    """
    if metadata is None:
        return None
    candidate: Any = None
    try:
        if isinstance(metadata, dict):
            candidate = metadata.get("Title")
        else:
            candidate = getattr(metadata, "title", None)
    except Exception:  # pragma: no cover -- defensive
        return None
    if isinstance(candidate, str) and candidate.strip():
        return candidate.strip()
    return None


def _render_markdown_table(
    table: list[list[Any]],
    *,
    page_index: int,
    url: str,
    max_cells: int,
) -> Optional[str]:
    """Render a pdfplumber table (list of rows) as a GFM markdown table.

    The first row becomes the header. ``None`` cells (very common in
    pdfplumber output) render as empty; pipes / newlines inside a cell
    are escaped so they don't break the table grid. Bounded by
    ``max_cells`` (rows * columns): when exceeded, whole trailing rows
    are dropped to stay under budget and a WARNING is logged. Returns
    ``None`` for an empty/degenerate table.
    """
    rows = [r for r in table if r is not None]
    if not rows:
        return None
    width = max(len(r) for r in rows)
    if width == 0:
        return None

    if max_cells > 0 and width * len(rows) > max_cells:
        keep_rows = max(1, max_cells // width)
        if keep_rows < len(rows):
            logger.warning(
                "PDF table on page {n} exceeds cell cap ({cap}); keeping "
                "{kept}/{total} rows for {url}",
                n=page_index,
                cap=max_cells,
                kept=keep_rows,
                total=len(rows),
                url=url,
            )
            rows = rows[:keep_rows]

    def _cell(value: Any) -> str:
        text = "" if value is None else str(value)
        # Collapse newlines (a cell spanning lines would break the row)
        # and escape pipes so the grid stays intact.
        return text.replace("\r", " ").replace("\n", " ").replace("|", "\\|").strip()

    def _row(cells: list[Any]) -> str:
        padded = list(cells) + [None] * (width - len(cells))
        return "| " + " | ".join(_cell(c) for c in padded) + " |"

    header = _row(rows[0])
    separator = "| " + " | ".join(["---"] * width) + " |"
    body = [_row(r) for r in rows[1:]]
    return "\n".join([header, separator, *body])


def _extract_json_ld(html: str) -> list[dict[str, Any]]:
    """v1.6.12: parse ``<script type='application/ld+json'>`` blocks.

    Returns a flat list of schema.org objects embedded in the page.
    ``@graph`` containers are unwrapped so the result contains
    individual items, not the graph wrapper. Malformed JSON-LD
    (sadly very common -- trailing commas, single-quoted keys, etc.)
    is swallowed silently; we never raise from this helper.

    Args:
        html: Raw page HTML.

    Returns:
        Flat list of dict objects (empty when no JSON-LD found / all
        blocks malformed).
    """
    if not html:
        return []
    # v1.6.14 E-3: cap total JSON-LD blocks so a hostile page can't force
    # unbounded growth of ``structured_data`` -- a single ``@graph`` with
    # 100k entries, or thousands of ld+json scripts, would otherwise be
    # accumulated wholesale into the result (and serialized downstream).
    # 500 is far above any legitimate page's structured-data count.
    max_blocks = 500
    blocks: list[dict[str, Any]] = []
    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:  # pragma: no cover -- defensive
        return []
    for script in soup.find_all("script", type="application/ld+json"):
        if len(blocks) >= max_blocks:
            break
        # ``script.string`` is None when the tag has children (e.g.
        # CDATA-wrapped); fall back to ``get_text()``.
        text = script.string if script.string else script.get_text()
        if not text or not text.strip():
            continue
        try:
            data = json.loads(text)
        except (json.JSONDecodeError, ValueError, RecursionError, MemoryError):
            # Many sites have malformed JSON-LD. Swallow and continue.
            # ``RecursionError`` covers an adversarial DoS: an attacker
            # could embed deeply-nested JSON (>1000 levels) to crash the
            # default CPython parser. ``MemoryError`` (v1.6.14 E-3) covers a
            # huge flat blob. Neither derives from ``ValueError`` /
            # ``JSONDecodeError``, so they must be caught explicitly here.
            continue
        # JSON-LD allows either a single object or an array at the
        # top level. Handle both.
        candidates: list[Any] = data if isinstance(data, list) else [data]
        for item in candidates:
            if len(blocks) >= max_blocks:
                break
            if not isinstance(item, dict):
                continue
            # Unwrap ``@graph`` containers -- the graph wrapper is
            # rarely what callers want; they want the contained items.
            graph = item.get("@graph")
            if isinstance(graph, list):
                # v1.6.14 E-3: bounded so a single giant @graph can't blow
                # past the cap in one extend().
                for g in graph:
                    if len(blocks) >= max_blocks:
                        break
                    if isinstance(g, dict):
                        blocks.append(g)
            else:
                blocks.append(item)
    return blocks


def slice_text_window(
    text: str,
    *,
    offset: int = 0,
    max_chars: Optional[int] = None,
) -> tuple[str, Optional[int]]:
    """v1.7.0: return a ``(window, next_offset)`` slice of ``text``.

    The continuation primitive behind ``ExtractionResult.truncated`` /
    ``next_offset``: ``window = text[offset : offset + max_chars]`` with a
    newline-boundary adjustment, and ``next_offset`` is the absolute offset
    to pass on the next call (``None`` when the window reaches the end of
    ``text``).

    Semantics:

    - ``offset`` is clamped to ``>= 0``. An offset at/past the end of
      ``text`` yields ``("", None)``.
    - ``max_chars=None`` means unlimited (window runs to the end).
    - ``max_chars`` is clamped to ``>= 1`` so a hostile/buggy value can
      never produce a zero-progress paging loop.
    - When the window would cut mid-text, the cut snaps back to the last
      newline found within the final 200 chars of the window (the newline
      is INCLUDED in the window) so chunks rejoin losslessly and the cut
      lands on a line boundary when one is near.
    """
    total = len(text)
    start = max(0, offset)
    if start >= total:
        return "", None
    if max_chars is None:
        return text[start:], None
    cap = max(1, max_chars)
    end = start + cap
    if end >= total:
        return text[start:], None
    nl = text.rfind("\n", max(start, end - 200), end)
    if nl != -1:
        end = nl + 1
    return text[start:end], end


def _synthesize_fetch_failure_message(status_value: str, status_code: Optional[int]) -> str:
    """v1.7.0: actionable error text when the fetcher supplied none.

    Keyed off the FetchStatus VALUE string and the HTTP status code so it
    stays valid for any status the enum grows. Every message names a next
    step the calling agent can take.
    """
    if status_code == 403:
        return "server returned 403; consider an authenticated session or a different source"
    if status_code == 404:
        return "server returned 404 (not found); the URL may be stale -- search for a fresh link"
    if status_code == 429:
        return "server returned 429 (rate limited); wait before retrying or reduce request rate"
    if status_code is not None:
        return f"server returned HTTP {status_code}; try a different source"
    if status_value == "timeout":
        return "page load timed out; retry once, or try a different source if it persists"
    if status_value == "blocked":
        return (
            "robots.txt or the domain allow/deny policy forbids fetching this path; "
            "do not retry this URL"
        )
    if status_value == "network_error":
        return "network error while fetching; verify the URL is reachable or try a different source"
    return f"fetch failed with status '{status_value}'; try a different source"


def build_fetch_failure_result(fetch_result: FetchResult) -> ExtractionResult:
    """v1.7.0: build a failure-transparent ExtractionResult from a failed fetch.

    Replaces the pre-v1.7.0 opaque ``ExtractionResult(url=...,
    extraction_method="none")`` for non-success fetches: the FetchResult's
    status / status_code / error_message are now CARRIED onto the
    extraction so callers (and LLMs behind MCP) can distinguish a 403
    bot-wall from a robots block from a timeout, and self-correct.

    The fetcher's ``error_message`` is carried verbatim when present;
    otherwise an actionable message is synthesized from the status value
    and HTTP code. ``failure_stage`` is always ``'fetch'`` here. Written
    generically off the status VALUE so any future FetchStatus member is
    propagated without changes here.
    """
    status_value = str(getattr(fetch_result.status, "value", fetch_result.status))
    message = (fetch_result.error_message or "").strip()
    if not message:
        message = _synthesize_fetch_failure_message(status_value, fetch_result.status_code)
    return ExtractionResult(
        url=fetch_result.url,
        extraction_method="none",
        fetch_status=status_value,
        status_code=fetch_result.status_code,
        error_message=message,
        failure_stage="fetch",
        correlation_id=fetch_result.correlation_id,
    )


def _apply_content_window(
    result: ExtractionResult,
    *,
    offset: int,
    max_chars: Optional[int],
) -> ExtractionResult:
    """v1.7.0: slice ``result``'s text to a window and stamp the metadata.

    The PRIMARY text is ``content`` when present, else ``markdown`` (in
    practice ``markdown`` never exists without ``content`` -- trafilatura
    populates both). Window metadata (``truncated`` /
    ``total_content_chars`` / ``content_offset`` / ``next_offset``)
    describes the primary text; when both representations exist,
    ``markdown`` receives the same (offset, max_chars) window applied to
    its own text so neither field smuggles the full document.

    Mutates ``result`` in place (consistent with ``_cap_content``) and
    returns it. ``total_content_chars`` is preserved when already set by
    a prior cap (it then reflects the larger pre-cap length).
    """
    primary_is_content = result.content is not None
    primary = result.content if primary_is_content else result.markdown
    if primary is None:
        result.content_offset = max(0, offset)
        return result
    start = max(0, offset)
    window, next_off = slice_text_window(primary, offset=start, max_chars=max_chars)
    if primary_is_content:
        result.content = window
        result.content_length = len(window)
        if result.markdown is not None:
            md_window, _ = slice_text_window(result.markdown, offset=start, max_chars=max_chars)
            result.markdown = md_window
    else:
        result.markdown = window
    if result.total_content_chars is None:
        result.total_content_chars = len(primary)
    result.truncated = result.truncated or next_off is not None
    result.content_offset = start
    result.next_offset = next_off
    return result


def _is_pdf(fr: FetchResult) -> bool:
    ct = (fr.content_type or "").lower()
    if "application/pdf" in ct:
        return True
    return urlparse(fr.final_url or fr.url).path.lower().endswith(".pdf")


def _is_xlsx(fr: FetchResult) -> bool:
    ct = (fr.content_type or "").lower()
    if "spreadsheetml" in ct or "openxmlformats-officedocument.spreadsheet" in ct:
        return True
    return urlparse(fr.final_url or fr.url).path.lower().endswith(".xlsx")


def _is_docx(fr: FetchResult) -> bool:
    ct = (fr.content_type or "").lower()
    if "wordprocessingml" in ct or "officedocument.wordprocessing" in ct:
        return True
    return urlparse(fr.final_url or fr.url).path.lower().endswith(".docx")


def _is_csv(fr: FetchResult) -> bool:
    ct = (fr.content_type or "").lower()
    if ct.startswith(("text/csv", "text/tab-separated-values")):
        return True
    return urlparse(fr.final_url or fr.url).path.lower().endswith((".csv", ".tsv"))


class ContentExtractor:
    """Extracts structured content from raw HTML using a layered fallback strategy."""

    def __init__(self, config: AppConfig) -> None:
        self._config = config

    def extract(
        self,
        fetch_result: FetchResult,
        *,
        strict: bool = False,
        prefer_api: bool = False,
        max_chars: Optional[int] = None,
        offset: int = 0,
    ) -> ExtractionResult:
        """Extract structured content from a FetchResult.

        Dispatches on FetchResult contents:
          - ``binary`` populated -> PDF or XLSX branch (requires ``[binary]`` extra)
          - ``html`` populated -> three-tier HTML fallback chain
            (+ v1.6.12 optional API-candidates path when ``prefer_api=True``)

        HTML fallback chain:
          1. (v1.6.12) ``prefer_api=True`` AND a captured XHR/fetch JSON
             response body is available -> ``api_json`` extraction.
          2. trafilatura (best quality, F1 ~0.958)
          3. BeautifulSoup4 structural extraction
          4. Raw text stripping (last resort)

        v1.6.12: every HTML result is enriched with
        ``structured_data`` -- parsed ``<script type='application/ld+
        json'>`` blocks -- always-on (cheap, no opt-in needed). When
        the page has no JSON-LD or all blocks are malformed, the list
        is empty.

        Args:
            fetch_result: The FetchResult to extract from.
            strict: If True, raise :class:`ExtractionError` when all
                three layers fail to produce content (very rare, since
                raw is a last-resort always-success path). When False,
                returns an ExtractionResult with extraction_method="none".
            prefer_api: v1.6.12. When True AND ``fetch_result.network_events``
                contains a response event with a captured JSON body
                (requires ``DiagnosticsConfig.capture_response_bodies=
                True``), route extraction through that body instead of
                the rendered HTML. Useful on SPAs where the XHR payload
                is strictly cleaner than the DOM. Default False
                preserves the v1.6.11 behaviour.
            max_chars: v1.7.0. Optional content window size in chars,
                applied AFTER extraction (and after the safety char cap).
                ``None`` (default) keeps the Python API unlimited -- the
                MCP boundary is where ``ExtractionConfig.default_max_chars``
                bites. When set, ``content`` (and ``markdown``, when
                present) are sliced and ``truncated`` /
                ``total_content_chars`` / ``content_offset`` /
                ``next_offset`` are populated. Slices snap back to a
                newline within the last 200 chars of the window.
            offset: v1.7.0. Continuation offset (chars into the full
                extracted text) for paging through large documents. Pass
                the previous result's ``next_offset``. An offset at/past
                the end yields empty content with ``next_offset=None``.

        Raises:
            ExtractionError: If ``strict=True`` and all layers fail.
        """
        result = self._extract_dispatch(fetch_result, strict=strict, prefer_api=prefer_api)
        # v1.6.16 CE-1: cap the emitted content for EVERY branch, not just
        # the api_json path (which got the E-6 cap at ``_extract_from_api_
        # candidates``). The binary (PDF/XLSX/DOCX/CSV) and HTML
        # (trafilatura/bs4/raw) extractors returned ``content=text`` with no
        # bound, so a hostile/huge document (input bounded only by
        # ``download.max_file_size_mb``, default 100 MB) could yield ~100 MB
        # of extracted text and blow up downstream token budgets / JSON
        # serialization. ``BudgetTracker.add_chars`` only raises after the
        # fact and isn't even called on the single-URL extract path, so it
        # does not save this. Truncate centrally here -- one cap for all
        # branches -- and keep ``content_length`` honest (mirrors the
        # api_json slice). The api_json branch's own 512 KiB cap is more
        # conservative and still applies first, so it is preserved.
        result = self._cap_content(result)
        # v1.7.0: optional content window (slicing) AFTER extraction + cap.
        # No-op on the default path (max_chars=None, offset=0) so existing
        # Python-API callers see byte-identical behaviour.
        if max_chars is not None or offset > 0:
            result = _apply_content_window(result, offset=offset, max_chars=max_chars)
        # v1.7.0 Wave 3A: strip invisible/bidi chars + run advisory injection
        # detection on the final VISIBLE text. Runs last so the scan sees
        # exactly what a caller would consume. Gated on the safety flags.
        result = self._apply_injection_containment(result)
        return result

    async def extract_async(
        self,
        fetch_result: FetchResult,
        *,
        strict: bool = False,
        prefer_api: bool = False,
        max_chars: Optional[int] = None,
        offset: int = 0,
    ) -> ExtractionResult:
        """Run :meth:`extract` off the event loop in a worker thread.

        v1.6.16 deep-review fix: ``extract`` drives synchronous CPU-heavy
        parsers (trafilatura, lxml, pypdf, openpyxl) over potentially multi-MB
        documents. Calling it directly from an async method blocks the event
        loop for the whole parse, stalling every other in-flight fetch/extract
        (especially harmful for the long-lived MCP server). The extractors build
        fresh per-call parser state, so running concurrent calls in threads is
        safe. Signature mirrors :meth:`extract`.
        """
        return await asyncio.to_thread(
            self.extract,
            fetch_result,
            strict=strict,
            prefer_api=prefer_api,
            max_chars=max_chars,
            offset=offset,
        )

    def _cap_content(self, result: ExtractionResult) -> ExtractionResult:
        """Truncate ``content``/``markdown`` to ``safety.max_chars_per_call``.

        Reuses the existing per-call character budget field rather than a
        new knob. Slices in place (consistent with the api_json E-6 cap)
        and resets ``content_length`` to the truncated length so it stays
        honest. v1.7.0: when the cap fires it now also sets ``truncated``
        and ``total_content_chars`` (with ``next_offset`` left ``None`` --
        the overflow is dropped, not pageable) so the cut is no longer
        silent.
        """
        cap = self._config.safety.max_chars_per_call
        if cap and result.content is not None and len(result.content) > cap:
            if result.total_content_chars is None:
                result.total_content_chars = len(result.content)
            result.content = result.content[:cap]
            result.content_length = len(result.content)
            result.truncated = True
        if cap and result.markdown is not None and len(result.markdown) > cap:
            if result.total_content_chars is None:
                result.total_content_chars = len(result.markdown)
            result.markdown = result.markdown[:cap]
            result.truncated = True
        return result

    def _apply_injection_containment(self, result: ExtractionResult) -> ExtractionResult:
        """v1.7.0 Wave 3A: strip invisible chars + run advisory injection scan.

        Composes with the rest of the pipeline (runs AFTER cap + window so it
        sees exactly the text a caller consumes). Two independent gates:

        - ``SafetyConfig.sanitize_fetched_content`` (default on): strip
          zero-width / bidi-control / invisible characters from ``content``
          and ``markdown``, accumulate the removed count onto the report, and
          set ``content_sanitized=True``. (Hidden-DOM stripping already ran
          pre-extraction; its count is carried on ``result.injection`` here.)
        - ``SafetyConfig.detect_prompt_injection`` (default on): run
          :func:`detect_injection` on the visible text and merge the risk /
          indicators / score into the report. On a HIGH detection the
          configured ``injection_action`` is applied (``flag`` = no change;
          ``redact`` = mask matched spans; ``block`` = empty the content and
          stamp an actionable ``error_message``).

        When BOTH gates are off this is a near no-op: any hidden-element count
        captured pre-extraction is preserved, but no new report is created
        and ``content_sanitized`` stays False.
        """
        safety = self._config.safety
        sanitize = safety.sanitize_fetched_content
        detect = safety.detect_prompt_injection

        # Preserve any hidden-element count stamped pre-extraction.
        existing = result.injection
        stripped_hidden = existing.stripped_hidden_elements if existing is not None else 0

        if not sanitize and not detect:
            return result

        invisible_removed = 0
        if sanitize:
            if result.content is not None:
                cleaned, n = strip_invisible_chars(result.content)
                if n:
                    result.content = cleaned
                    result.content_length = len(cleaned)
                    invisible_removed += n
            if result.markdown is not None:
                cleaned_md, n = strip_invisible_chars(result.markdown)
                if n:
                    result.markdown = cleaned_md
                    invisible_removed += n
            result.content_sanitized = True

        if detect:
            visible = result.content if result.content is not None else (result.markdown or "")
            report = detect_injection(visible)
            report.stripped_hidden_elements = stripped_hidden
            report.stripped_invisible_chars = invisible_removed
            result.injection = report
            if report.risk == "high":
                result = self._apply_injection_action(result, report)
        elif sanitize and (stripped_hidden or invisible_removed):
            # Detection off but sanitization ran and removed something --
            # keep a counts-only report so the strip is observable.
            result.injection = InjectionReport(
                stripped_hidden_elements=stripped_hidden,
                stripped_invisible_chars=invisible_removed,
            )

        return result

    def _apply_injection_action(
        self, result: ExtractionResult, report: InjectionReport
    ) -> ExtractionResult:
        """Apply ``SafetyConfig.injection_action`` to a HIGH-risk result.

        ``flag`` (default) leaves content untouched -- the report alone is the
        signal. ``redact`` masks the matched indicator phrases in the visible
        text. ``block`` empties the content and stamps an actionable
        ``error_message`` (explicit opt-in only).
        """
        action = self._config.safety.injection_action
        if action == "flag":
            return result
        if action == "redact":
            # v1.7.0 (gap-fix): mask real match spans via redact_injection
            # rather than reconstructing a needle from the truncated /
            # whitespace-collapsed indicator snippet (which rarely appears
            # verbatim, making the old loop a silent no-op).
            for field in ("content", "markdown"):
                text = getattr(result, field)
                if not text:
                    continue
                redacted, _ = redact_injection(text)
                setattr(result, field, redacted)
            if result.content is not None:
                result.content_length = len(result.content)
            return result
        # action == "block"
        result.content = None
        result.markdown = None
        result.content_length = 0
        result.error_message = (
            "content blocked: high-confidence prompt-injection indicators "
            f"detected (risk={report.risk}, score={report.score}); "
            "SafetyConfig.injection_action='block'. Inspect "
            "injection.indicators to see why, or set injection_action='flag' "
            "to receive the content with an advisory flag instead"
        )
        result.failure_stage = "injection_blocked"
        return result

    def _extract_dispatch(
        self,
        fetch_result: FetchResult,
        *,
        strict: bool = False,
        prefer_api: bool = False,
    ) -> ExtractionResult:
        """Dispatch to the binary / HTML extractors (uncapped).

        Split out from :meth:`extract` so the public method can apply a
        single content cap (CE-1) over every return path. See
        :meth:`extract` for the full contract.
        """
        if fetch_result.status != FetchStatus.SUCCESS:
            if strict:
                from .exceptions import ExtractionError

                raise ExtractionError(
                    f"Cannot extract from non-success FetchResult: "
                    f"status={fetch_result.status}, url={fetch_result.url}"
                )
            # v1.7.0 failure transparency: carry status / status_code /
            # error_message (verbatim, or synthesized actionably) onto the
            # ExtractionResult instead of returning opaque emptiness.
            return build_fetch_failure_result(fetch_result)

        url = fetch_result.final_url

        # Binary branch: PDF / XLSX / DOCX / CSV. Dispatched on
        # content_type or URL extension.
        if fetch_result.binary is not None:
            if _is_pdf(fetch_result):
                return self._extract_pdf(fetch_result.binary, url)
            if _is_xlsx(fetch_result):
                return self._extract_xlsx(fetch_result.binary, url)
            if _is_docx(fetch_result):
                return self._extract_docx(fetch_result.binary, url)
            if _is_csv(fetch_result):
                return self._extract_csv(fetch_result.binary, url)
            # Unrecognized binary -- not extractable. v1.7.0: say so.
            return ExtractionResult(
                url=url,
                extraction_method="none",
                content=None,
                error_message=(
                    f"binary content (content_type="
                    f"{fetch_result.content_type or 'unknown'}) is not an "
                    "extractable kind (pdf/xlsx/docx/csv); use web_download "
                    "to save the file instead"
                ),
            )

        if not fetch_result.html:
            if strict:
                from .exceptions import ExtractionError

                raise ExtractionError(f"FetchResult has neither html nor binary: url={url}")
            # v1.7.0: a SUCCESS fetch with no payload is a capture-level
            # emptiness -- say so instead of returning bare "none".
            return ExtractionResult(
                url=url,
                extraction_method="none",
                fetch_status=str(getattr(fetch_result.status, "value", fetch_result.status)),
                status_code=fetch_result.status_code,
                error_message=(
                    "fetch succeeded but returned no HTML content (empty page "
                    "or capture failure); retry once, or inspect the page with "
                    "web_observe / web_screenshot"
                ),
                failure_stage="fetch",
            )

        html = fetch_result.html
        min_len = self._config.extraction.min_content_length

        # v1.7.0 Wave 3A: strip hidden-from-humans DOM BEFORE extraction so
        # the injected text a human can't see (display:none divs, off-screen
        # spans, aria-hidden, comments, script/style) never reaches
        # trafilatura/bs4/markdown. Deterministic, zero false-positive harm.
        # The removed-element count is stashed on the result via an early
        # InjectionReport so ``extract`` can complete the report after the
        # invisible-char pass. Gated on the sanitize flag.
        # Keep the original HTML for JSON-LD extraction below: strip_hidden_dom
        # removes ALL <script> tags, which would also drop
        # <script type="application/ld+json"> metadata. JSON-LD is structured
        # DATA the caller wants, not hidden injection TEXT, so it is read from
        # the pre-strip HTML (it is surfaced in structured_data, never executed,
        # and the main-content injection scan still runs on the stripped text).
        original_html = html
        stripped_hidden = 0
        if self._config.safety.sanitize_fetched_content:
            html, stripped_hidden = strip_hidden_dom(html)

        def _stamp_hidden(res: ExtractionResult) -> ExtractionResult:
            if stripped_hidden:
                res.injection = InjectionReport(stripped_hidden_elements=stripped_hidden)
            return res

        # v1.6.12: JSON-LD enrichment. The helper swallows malformed
        # JSON (including ``RecursionError`` from deeply-nested
        # adversarial payloads). Cost: one extra BS4+lxml parse per
        # fetch (~5-20 ms on a 50-200 KB page). On the bs4 / raw
        # fallback paths this duplicates the parse those extractors
        # already do; a future patch can share the parsed soup, but
        # the duplication is acceptable for now (small absolute cost,
        # keeps the JSON-LD path decoupled from the fallback chain).
        # Compute once; attach to whichever extractor wins. Read from the
        # pre-strip HTML so hidden-DOM stripping doesn't drop JSON-LD scripts.
        structured = _extract_json_ld(original_html)

        # v1.6.12: prefer_api path -- when the caller opted in AND the
        # FetchResult carries a captured JSON response body, route
        # extraction through that body instead of the rendered HTML.
        if prefer_api:
            api_result = self._extract_from_api_candidates(fetch_result, url)
            if api_result is not None:
                api_result.structured_data = structured
                return api_result
            logger.debug(
                "prefer_api=True but no usable JSON body captured for {url}; "
                "falling back to HTML extraction",
                url=url,
            )

        # Layer 1: trafilatura
        result = self._extract_trafilatura(html, url)
        if result and result.content and len(result.content) >= min_len:
            result.structured_data = structured
            return _stamp_hidden(result)
        logger.debug("Trafilatura insufficient for {url}, trying BS4", url=url)

        # Layer 2: BeautifulSoup
        result = self._extract_bs4(html, url)
        if result and result.content and len(result.content) >= min_len:
            result.structured_data = structured
            return _stamp_hidden(result)
        logger.debug("BS4 insufficient for {url}, falling back to raw", url=url)

        # Layer 3: raw text (always-success unless catastrophic)
        raw_result = self._extract_raw(html, url)
        if strict and (not raw_result.content or len(raw_result.content) < min_len):
            from .exceptions import ExtractionError

            raise ExtractionError(
                f"All three extraction layers failed for {url} "
                f"(content_length={raw_result.content_length})"
            )
        raw_result.structured_data = structured
        return _stamp_hidden(raw_result)

    # ------------------------------------------------------------------
    # v1.6.12: API-candidates extraction (prefer_api=True path)
    # ------------------------------------------------------------------

    def _extract_from_api_candidates(
        self, fetch_result: FetchResult, url: str
    ) -> Optional[ExtractionResult]:
        """v1.6.12: extract from a captured XHR/fetch JSON response body.

        Scans :attr:`FetchResult.network_events` for response events
        with a captured ``body_text`` (requires
        ``DiagnosticsConfig.capture_response_bodies=True``) and JSON
        content-type. Among those it SCORES candidates (v1.6.16 CE-3:
        same-origin preferred, analytics/telemetry URLs deprioritised,
        JSON object/array root preferred; body size is only the
        tie-breaker) and uses the best as the extraction source. Returns
        ``None`` when no candidate found -- caller falls back to HTML.

        The extracted ``content`` is the pretty-printed JSON; ``title``
        is derived heuristically from common top-level keys
        (``title`` / ``headline`` / ``name``) when present. The parsed
        JSON is NOT put into ``structured_data`` -- that field is
        reserved for JSON-LD blocks parsed from the rendered HTML;
        ``extract`` populates it separately.

        Heuristic limitations (known, documented for caller awareness):

        - **Heuristic relevance, not exact.** v1.6.16 CE-3 scores
          candidates (same-origin + JSON-root preferred, analytics/
          telemetry URLs deprioritised) so a large Segment/GA4 batched-hit
          payload no longer beats a small same-origin API response -- but
          it is still a heuristic. Watch the DEBUG log line (it prints the
          chosen URL + score) to spot a mis-pick.
        - **No URL-pattern filter.** Cannot tell the caller "use the
          response from ``/api/page-data``, not ``/v1/track``". A
          future patch can accept a regex / glob.
        - **Title heuristic is top-level only.** Nested shapes like
          ``{data: {title: "..."}}`` won't populate
          ``ExtractionResult.title``, though the JSON body still ends
          up in ``content`` and a caller can parse it themselves.
        """
        if not fetch_result.network_events:
            return None
        # v1.6.16 CE-3: score candidates rather than blindly taking the
        # largest body. Score = same-origin as the page (+2) - analytics/
        # telemetry-looking URL (3) + JSON object/array root (+1); the body
        # size is only the tie-breaker. This stops a large batched analytics
        # payload from being picked over a small, relevant API response.
        page_host = (urlparse(url).hostname or "").lower()
        best_evt = None
        best_key: tuple[int, int] = (-(10**9), -1)
        for evt in fetch_result.network_events:
            if evt.event_type != "response":
                continue
            if evt.resource_type not in {"xhr", "fetch"}:
                continue
            if not evt.body_text:
                continue
            ct = (evt.content_type or "").lower()
            if "json" not in ct:
                continue
            evt_host = (urlparse(evt.url or "").hostname or "").lower()
            score = 0
            if page_host and evt_host == page_host:
                score += 2
            if _ANALYTICS_URL_RE.search(evt.url or ""):
                score -= 3
            if evt.body_text.lstrip()[:1] in ("{", "["):
                score += 1
            cand_key = (score, len(evt.body_text))
            if cand_key > best_key:
                best_key = cand_key
                best_evt = evt
        if best_evt is None or not best_evt.body_text:
            return None
        best_size = len(best_evt.body_text)
        try:
            parsed = json.loads(best_evt.body_text)
        except (json.JSONDecodeError, ValueError, RecursionError):
            return None
        # v1.6.12: emit the chosen URL at DEBUG so callers spotting
        # garbage extractions (analytics ping picked over real API)
        # have a diagnostic signal. The heuristic is documented in
        # the docstring; this log line is the runtime telemetry.
        logger.debug(
            "prefer_api selected XHR/fetch body for {url}: chosen_url={chosen} "
            "body_size={size} score={score}",
            url=url,
            chosen=best_evt.url,
            size=best_size,
            score=best_key[0],
        )

        # Heuristic title from top-level keys (defensive against None /
        # non-string types).
        title: Optional[str] = None
        if isinstance(parsed, dict):
            for key in ("title", "headline", "name"):
                val = parsed.get(key)
                if isinstance(val, str) and val.strip():
                    title = val.strip()
                    break

        # Pretty-print as the content payload.
        try:
            pretty = json.dumps(parsed, indent=2, ensure_ascii=False)
        except (TypeError, ValueError):
            pretty = best_evt.body_text  # fall back to raw

        # v1.6.14 E-6: cap the emitted content. The capture-time body cap
        # bounds ``best_evt.body_text``, NOT this re-serialized ``pretty``
        # output: indent=2 re-formatting can inflate a compact JSON body
        # several-fold (a 256 KiB compact blob with many short keys can
        # balloon past 1 MB), so the prior "truncation is implicit" claim
        # was wrong. Hard-cap here so content_length stays honest and a
        # hostile API body can't blow up downstream token budgets.
        max_api_content = 512 * 1024
        if pretty and len(pretty) > max_api_content:
            pretty = pretty[:max_api_content]

        return ExtractionResult(
            url=url,
            title=title,
            content=pretty,
            extraction_method="api_json",
            content_length=len(pretty),
        )

    # ------------------------------------------------------------------
    # Binary branches: PDF + XLSX
    # ------------------------------------------------------------------

    def _extract_pdf(self, blob: bytes, url: str) -> ExtractionResult:
        """Extract text (and tables) from PDF bytes.

        Engine preference chain (v1.7.0 Wave 3C):

        1. **pdfplumber** when installed -- per-page text with better
           layout fidelity AND per-page table extraction rendered as
           markdown. ``extraction_method='pdfplumber'``.
        2. **pypdf** fallback when pdfplumber is missing -- page text
           only, no tables. ``extraction_method='pdf'``.
        3. Neither installed -> ``extraction_method='none'`` with an
           install hint (the pre-Wave-3C missing-extra contract).

        Both engines emit per-page markers (``===== Page N =====``) when
        ``ExtractionConfig.pdf_page_markers`` is on so an agent can cite
        pages, and set ``page_count``. When a PDF has pages but no
        embedded text layer (scanned / image-only), the result is NOT a
        bare empty success -- ``error_message`` names the OCR reason and
        ``failure_stage='extract'``.

        Availability is probed with :func:`importlib.util.find_spec`
        (mirrors how the optional-extra paths detect their library
        without a hard top-level import) so the no-extra install keeps
        working; tests patch ``_module_available`` to exercise each branch.
        """
        if _module_available("pdfplumber"):
            return self._extract_pdf_pdfplumber(blob, url)
        if _module_available("pypdf"):
            return self._extract_pdf_pypdf(blob, url)
        logger.warning(
            "Neither pdfplumber nor pypdf installed; PDF extraction skipped. "
            "Install with: pip install 'web-agent-toolkit[binary]'"
        )
        return ExtractionResult(
            url=url,
            extraction_method="none",
            content=None,
            error_message=(
                "PDF extraction requires the [binary] extra (pdfplumber "
                "preferred, pypdf fallback); neither is installed. Install "
                "with: pip install 'web-agent-toolkit[binary]'"
            ),
            failure_stage="extract",
        )

    @staticmethod
    def _scanned_pdf_result(url: str, page_count: int) -> ExtractionResult:
        """Build the actionable result for an image-only / scanned PDF.

        A PDF with pages but no embedded text layer needs OCR, which
        web_agent does not perform. Rather than a bare empty success,
        surface why and what to do next.
        """
        logger.info(
            "PDF appears image-only / scanned (no text layer) for {url} "
            "({n} pages)",
            url=url,
            n=page_count,
        )
        return ExtractionResult(
            url=url,
            extraction_method="none",
            content=None,
            page_count=page_count,
            error_message=(
                "PDF appears to be image-only / scanned (no embedded text "
                "layer); OCR is required and is not supported by web_agent -- "
                "try a different source or an OCR tool."
            ),
            failure_stage="extract",
        )

    def _extract_pdf_pdfplumber(self, blob: bytes, url: str) -> ExtractionResult:
        """Preferred PDF path: per-page text + markdown tables via pdfplumber.

        Bounds table extraction by ``pdf_max_tables`` (total tables) and
        ``pdf_max_table_cells`` (per-table cells); both truncate with a
        WARNING and set ``truncated`` rather than silently dropping. The
        document is streamed page-by-page; only the rendered strings are
        retained, keeping peak memory bounded by the content cap rather
        than the raw page objects.
        """
        from io import BytesIO

        import pdfplumber

        ext_cfg = self._config.extraction
        emit_markers = ext_cfg.pdf_page_markers
        want_tables = ext_cfg.pdf_extract_tables and ext_cfg.pdf_max_tables > 0
        max_tables = ext_cfg.pdf_max_tables
        max_cells = ext_cfg.pdf_max_table_cells

        page_parts: list[str] = []
        tables_md: list[str] = []
        any_text = False
        tables_truncated = False
        title: Optional[str] = None
        try:
            with pdfplumber.open(BytesIO(blob)) as pdf:
                title = _pdf_title_from_metadata(getattr(pdf, "metadata", None))
                page_count = len(pdf.pages)
                for index, page in enumerate(pdf.pages, start=1):
                    segments: list[str] = []
                    if emit_markers:
                        segments.append(f"===== Page {index} =====")
                    try:
                        page_text = page.extract_text() or ""
                    except Exception as page_exc:
                        logger.debug(
                            "pdfplumber page {n} text extract failed: {e}",
                            n=index,
                            e=page_exc,
                        )
                        page_text = ""
                    if page_text.strip():
                        any_text = True
                        segments.append(page_text.strip())
                    if want_tables:
                        remaining = max_tables - len(tables_md)
                        page_tables, dropped = self._render_pdf_tables(
                            page,
                            page_index=index,
                            url=url,
                            remaining=remaining,
                            max_cells=max_cells,
                        )
                        if dropped:
                            tables_truncated = True
                        for md in page_tables:
                            tables_md.append(md)
                            segments.append(md)
                    if len(segments) > (1 if emit_markers else 0):
                        page_parts.append("\n\n".join(segments))
                    elif emit_markers:
                        # Keep the marker even for an empty page so page
                        # numbers stay aligned for citation.
                        page_parts.append(segments[0])
        except Exception as exc:
            logger.warning(
                "pdfplumber extraction failed for {url}: {e}", url=url, e=exc
            )
            return ExtractionResult(url=url, extraction_method="none", content=None)

        if page_count == 0:
            return ExtractionResult(url=url, extraction_method="none", content=None)
        if not any_text and not tables_md:
            # Pages exist but carry no extractable text layer or tables.
            return self._scanned_pdf_result(url, page_count)

        text = "\n\n".join(page_parts).strip()
        result = ExtractionResult(
            url=url,
            title=title,
            content=text if text else None,
            extraction_method="pdfplumber",
            content_length=len(text),
            page_count=page_count,
            tables=tables_md,
        )
        if tables_truncated:
            result.truncated = True
        return result

    def _render_pdf_tables(
        self,
        page: Any,
        *,
        page_index: int,
        url: str,
        remaining: int,
        max_cells: int,
    ) -> tuple[list[str], int]:
        """Extract a page's tables and render each as a markdown table.

        ``remaining`` is the global per-document table budget still
        available (``pdf_max_tables`` minus tables already rendered).
        Renders at most ``remaining`` tables from this page; returns
        ``(markdown_tables, dropped)`` where ``dropped`` is the number of
        this page's tables that were NOT rendered because the budget ran
        out (so the caller can flag ``truncated``). Per-table cell count
        is bounded by ``max_cells`` (whole trailing rows dropped) with a
        WARNING; an over-budget table is still emitted, just truncated.
        """
        try:
            raw_tables = page.extract_tables() or []
        except Exception as exc:
            logger.debug(
                "pdfplumber table extract failed on page {n}: {e}",
                n=page_index,
                e=exc,
            )
            return [], 0

        # Render only non-degenerate tables, capped by the remaining budget.
        renderable = [
            md
            for md in (
                _render_markdown_table(
                    t, page_index=page_index, url=url, max_cells=max_cells
                )
                for t in raw_tables
            )
            if md
        ]
        budget = max(0, remaining)
        if len(renderable) > budget:
            dropped = len(renderable) - budget
            logger.warning(
                "PDF table cap reached for {url}; dropping {dropped} table(s) "
                "(page {n})",
                url=url,
                dropped=dropped,
                n=page_index,
            )
            return renderable[:budget], dropped
        return renderable, 0

    def _extract_pdf_pypdf(self, blob: bytes, url: str) -> ExtractionResult:
        """Fallback PDF path: page text via pypdf (no tables).

        Adds per-page markers (when enabled) and ``page_count`` so the
        fallback still supports page citation, and the same scanned /
        no-text-layer detection as the pdfplumber path. ``extraction_
        method='pdf'`` so callers/telemetry can see pypdf won (vs
        ``'pdfplumber'``).
        """
        from io import BytesIO

        from pypdf import PdfReader

        emit_markers = self._config.extraction.pdf_page_markers
        try:
            reader = PdfReader(BytesIO(blob))
            if reader.is_encrypted:
                logger.info("PDF is encrypted, skipping: {url}", url=url)
                return ExtractionResult(url=url, extraction_method="none", content=None)
            page_parts: list[str] = []
            any_text = False
            page_count = len(reader.pages)
            for index, page in enumerate(reader.pages, start=1):
                try:
                    page_text = (page.extract_text() or "").strip()
                except Exception as page_exc:
                    logger.debug("PDF page extract failed: {e}", e=page_exc)
                    page_text = ""
                if page_text:
                    any_text = True
                if emit_markers:
                    marker = f"===== Page {index} ====="
                    page_parts.append(f"{marker}\n\n{page_text}" if page_text else marker)
                elif page_text:
                    page_parts.append(page_text)
            if page_count == 0:
                return ExtractionResult(url=url, extraction_method="none", content=None)
            if not any_text:
                return self._scanned_pdf_result(url, page_count)
            text = "\n\n".join(page_parts).strip()
            title = _pdf_title_from_metadata(reader.metadata)
            return ExtractionResult(
                url=url,
                title=title,
                content=text if text else None,
                extraction_method="pdf",
                content_length=len(text),
                page_count=page_count,
            )
        except Exception as exc:
            logger.warning("PDF extraction failed for {url}: {e}", url=url, e=exc)
            return ExtractionResult(url=url, extraction_method="none", content=None)

    @staticmethod
    def _extract_xlsx(blob: bytes, url: str) -> ExtractionResult:
        """Extract text from XLSX bytes using openpyxl (TSV-style per sheet)."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning(
                "openpyxl not installed; XLSX extraction skipped. "
                "Install with: pip install 'web-agent-toolkit[binary]'"
            )
            return ExtractionResult(
                url=url,
                extraction_method="none",
                content=None,
            )

        try:
            from io import BytesIO

            wb = load_workbook(BytesIO(blob), read_only=True, data_only=True)
            sheet_dumps: list[str] = []
            for sheet in wb.worksheets:
                rows: list[str] = [f"# Sheet: {sheet.title}"]
                for row in sheet.iter_rows(values_only=True):
                    cells = ["" if cell is None else str(cell) for cell in row]
                    if any(cells):
                        rows.append("\t".join(cells))
                if len(rows) > 1:
                    sheet_dumps.append("\n".join(rows))
            wb.close()
            text = "\n\n".join(sheet_dumps).strip()
            if not text:
                return ExtractionResult(
                    url=url,
                    extraction_method="none",
                    content=None,
                )
            return ExtractionResult(
                url=url,
                content=text,
                extraction_method="xlsx",
                content_length=len(text),
            )
        except Exception as exc:
            logger.warning("XLSX extraction failed for {url}: {e}", url=url, e=exc)
            return ExtractionResult(url=url, extraction_method="none", content=None)

    @staticmethod
    def _extract_docx(blob: bytes, url: str) -> ExtractionResult:
        """Extract text from DOCX bytes using python-docx (paragraph-by-paragraph)."""
        try:
            import docx as python_docx  # python-docx package
        except ImportError:
            logger.warning(
                "python-docx not installed; DOCX extraction skipped. "
                "Install with: pip install 'web-agent-toolkit[binary]'"
            )
            return ExtractionResult(url=url, extraction_method="none", content=None)

        try:
            from io import BytesIO

            doc = python_docx.Document(BytesIO(blob))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text:
                    parts.append(para.text)
            # Tables: dump cells row-by-row, tab-separated, with a marker.
            for table in doc.tables:
                parts.append("# Table")
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append("\t".join(cells))
            text = "\n".join(parts).strip()
            if not text:
                return ExtractionResult(url=url, extraction_method="none", content=None)
            # /Title from core properties when available
            title: Optional[str] = None
            try:
                cp = doc.core_properties
                if cp is not None and cp.title:
                    title = str(cp.title)
            except Exception:
                pass
            return ExtractionResult(
                url=url,
                title=title,
                content=text,
                extraction_method="docx",
                content_length=len(text),
            )
        except Exception as exc:
            logger.warning("DOCX extraction failed for {url}: {e}", url=url, e=exc)
            return ExtractionResult(url=url, extraction_method="none", content=None)

    @staticmethod
    def _extract_csv(blob: bytes, url: str) -> ExtractionResult:
        """Extract text from CSV/TSV bytes using stdlib csv (no dep required).

        Auto-detects delimiter via :class:`csv.Sniffer`. Falls back to comma
        on detection failure. Returns the raw content as TSV-style text so
        downstream LLM consumers see one row per line with tab separators.
        """
        try:
            import csv as csv_mod
            from io import StringIO

            # Decode bytes -> text. Try utf-8 first, then latin-1 as a
            # last-resort fallback. We never crash on encoding issues.
            text_in: str
            for enc in ("utf-8-sig", "utf-8", "latin-1"):
                try:
                    text_in = blob.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return ExtractionResult(url=url, extraction_method="none", content=None)

            sample = text_in[:4096]
            try:
                dialect = csv_mod.Sniffer().sniff(sample, delimiters=",;\t|")
            except Exception:
                dialect = csv_mod.excel  # default: comma-separated
            reader = csv_mod.reader(StringIO(text_in), dialect)
            rows = ["\t".join(row) for row in reader if any(c.strip() for c in row)]
            text = "\n".join(rows).strip()
            if not text:
                return ExtractionResult(url=url, extraction_method="none", content=None)
            return ExtractionResult(
                url=url,
                content=text,
                extraction_method="csv",
                content_length=len(text),
            )
        except Exception as exc:
            logger.warning("CSV extraction failed for {url}: {e}", url=url, e=exc)
            return ExtractionResult(url=url, extraction_method="none", content=None)

    def _extract_trafilatura(self, html: str, url: str) -> Optional[ExtractionResult]:
        """Primary extractor using trafilatura with metadata."""
        try:
            doc = trafilatura.bare_extraction(
                html,
                url=url,
                favor_precision=self._config.extraction.favor_precision,
                favor_recall=self._config.extraction.favor_recall,
                include_tables=self._config.extraction.include_tables,
                include_links=self._config.extraction.include_links,
                include_comments=self._config.extraction.include_comments,
                with_metadata=True,
            )
            if doc is None:
                return None

            # bare_extraction returns a Document object; access attributes directly
            text = getattr(doc, "text", None)
            if not text:
                return None

            # Second pass: ask trafilatura for a markdown rendering of
            # the same page. Cheap (HTML re-parsed once) and the result
            # is what most LLMs prefer to consume because it preserves
            # headings, lists, links, and emphasis. Best-effort -- on
            # failure we leave markdown=None.
            markdown: Optional[str] = None
            try:
                markdown = trafilatura.extract(
                    html,
                    url=url,
                    output_format="markdown",
                    favor_precision=self._config.extraction.favor_precision,
                    favor_recall=self._config.extraction.favor_recall,
                    include_tables=self._config.extraction.include_tables,
                    include_links=self._config.extraction.include_links,
                    include_comments=self._config.extraction.include_comments,
                )
            except Exception as md_exc:
                logger.debug("Markdown rendering failed for {url}: {e}", url=url, e=md_exc)

            return ExtractionResult(
                url=url,
                title=getattr(doc, "title", None),
                description=getattr(doc, "description", None),
                author=getattr(doc, "author", None),
                date=getattr(doc, "date", None),
                sitename=getattr(doc, "sitename", None),
                content=text,
                markdown=markdown,
                language=getattr(doc, "language", None),
                extraction_method="trafilatura",
                content_length=len(text),
            )
        except Exception as e:
            logger.warning("Trafilatura failed for {url}: {e}", url=url, e=e)
            return None

    def _extract_bs4(self, html: str, url: str) -> Optional[ExtractionResult]:
        """Fallback extractor using BeautifulSoup structural heuristics."""
        try:
            soup = BeautifulSoup(html, "lxml")

            # Title
            title = None
            title_tag = soup.find("title")
            if title_tag:
                title = title_tag.get_text(strip=True)

            # Meta description
            description: Optional[str] = None
            meta_desc = soup.find("meta", attrs={"name": "description"})
            if meta_desc:
                # bs4 >= 4.13 types Tag.get() as str | AttributeValueList | None;
                # ExtractionResult.description requires str | None, so we coerce.
                desc_val = meta_desc.get("content")
                description = str(desc_val) if desc_val is not None else None

            # Author
            author: Optional[str] = None
            meta_author = soup.find("meta", attrs={"name": "author"})
            if meta_author:
                author_val = meta_author.get("content")
                author = str(author_val) if author_val is not None else None

            # Main content: try semantic tags first, then common class/id patterns
            content_tag = (
                soup.find("article")
                or soup.find("main")
                or soup.find("div", {"role": "main"})
                or soup.find("div", class_="content")
                or soup.find("div", id="content")
                or soup.body
            )

            # Strip non-content elements
            if content_tag:
                for unwanted in content_tag.find_all(
                    ["nav", "header", "footer", "aside", "script", "style", "noscript"]
                ):
                    unwanted.decompose()
                text = content_tag.get_text(separator="\n", strip=True)
            else:
                text = ""

            if not text:
                return None

            return ExtractionResult(
                url=url,
                title=title,
                description=description,
                author=author,
                content=text,
                extraction_method="bs4",
                content_length=len(text),
            )
        except Exception as e:
            logger.warning("BS4 extraction failed for {url}: {e}", url=url, e=e)
            return None

    def _extract_raw(self, html: str, url: str) -> ExtractionResult:
        """Last resort: strip all tags and return body text."""
        try:
            soup = BeautifulSoup(html, "lxml")
            for tag in soup.find_all(["script", "style", "noscript"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
        except Exception:
            text = ""

        return ExtractionResult(
            url=url,
            content=text if text else None,
            extraction_method="raw",
            content_length=len(text) if text else 0,
        )
